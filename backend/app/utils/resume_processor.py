import os
import uuid
from typing import Optional, Dict, Any
from pathlib import Path
import re

# PDF processing
import PyPDF2
import pdfplumber

# DOCX processing
from docx import Document

# Text cleaning
import nltk
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords
import spacy

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

# Load spaCy model
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    print("spaCy English model not found. Please install it with: python -m spacy download en_core_web_sm")
    nlp = None

class ResumeProcessor:
    def __init__(self):
        self.stop_words = set(stopwords.words('english'))
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        
        try:
            # Try with pdfplumber first (better for structured PDFs)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"pdfplumber failed: {e}")
            
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            except Exception as e:
                print(f"PyPDF2 also failed: {e}")
                raise Exception(f"Failed to extract text from PDF: {e}")
        
        return text.strip()
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
            
            return "\n".join(text).strip()
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {e}")
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read().strip()
            except Exception as e:
                raise Exception(f"Failed to extract text from TXT: {e}")
        except Exception as e:
            raise Exception(f"Failed to extract text from TXT: {e}")
    
    def extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text based on file type"""
        file_type = file_type.lower()
        
        if file_type == 'pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_type == 'docx':
            return self.extract_text_from_docx(file_path)
        elif file_type == 'txt':
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove extra whitespace and newlines
        text = re.sub(r'\n+', ' ', text)
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep important punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)]', ' ', text)
        
        # Remove extra spaces
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()
    
    def extract_contact_info(self, text: str) -> Dict[str, Any]:
        """Extract contact information from resume text"""
        contact_info = {}
        
        # Email pattern
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            contact_info['email'] = emails[0]
        
        # Phone pattern (various formats)
        phone_patterns = [
            r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',  # US format
            r'\b\(\d{3}\)\s?\d{3}[-.\s]?\d{4}\b',  # (XXX) XXX-XXXX
            r'\b\+\d{1,3}[-.\s]?\d{3,14}\b'       # International
        ]
        
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                contact_info['phone'] = phones[0]
                break
        
        # LinkedIn profile
        linkedin_pattern = r'linkedin\.com/in/[\w\-]+'
        linkedin = re.findall(linkedin_pattern, text, re.IGNORECASE)
        if linkedin:
            contact_info['linkedin'] = f"https://{linkedin[0]}"
        
        return contact_info
    
    def extract_education(self, text: str) -> Dict[str, Any]:
        """Extract education information"""
        education = {
            'degrees': [],
            'institutions': [],
            'graduation_years': []
        }
        
        # Degree patterns
        degree_patterns = [
            r'\b(?:Bachelor|B\.A\.|B\.S\.|B\.Sc\.|BA|BS|BSc)\b.*?\b(?:in|of)\s+([A-Za-z\s]+)',
            r'\b(?:Master|M\.A\.|M\.S\.|M\.Sc\.|MA|MS|MSc|MBA)\b.*?\b(?:in|of)\s+([A-Za-z\s]+)',
            r'\b(?:Doctor|Ph\.D\.|PhD|doctorate)\b.*?\b(?:in|of)\s+([A-Za-z\s]+)',
            r'\b(?:Associate|A\.A\.|A\.S\.|AA|AS)\b.*?\b(?:in|of)\s+([A-Za-z\s]+)'
        ]
        
        for pattern in degree_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            education['degrees'].extend(matches)
        
        # University/College patterns
        institution_patterns = [
            r'\b(?:University|College|Institute|School)\s+(?:of\s+)?([A-Za-z\s]+)',
            r'\b([A-Za-z\s]+)\s+(?:University|College|Institute)'
        ]
        
        for pattern in institution_patterns:
            matches = re.findall(pattern, text)
            education['institutions'].extend(matches)
        
        # Graduation years
        year_pattern = r'\b(19|20)\d{2}\b'
        years = re.findall(year_pattern, text)
        education['graduation_years'] = list(set([''.join(year) for year in years]))
        
        return education
    
    def extract_experience(self, text: str) -> Dict[str, Any]:
        """Extract work experience information"""
        experience = {
            'total_years': 0,
            'companies': [],
            'positions': [],
            'experience_sections': []
        }
        
        # Common experience section headers
        experience_headers = [
            r'(?:work\s+)?experience',
            r'employment\s+history',
            r'professional\s+experience',
            r'career\s+history',
            r'work\s+history'
        ]
        
        # Try to find experience sections
        for header in experience_headers:
            pattern = f'{header}(.*?)(?=education|skills|projects|$)'
            matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
            if matches:
                experience['experience_sections'].extend(matches)
        
        # Extract company names (basic patterns)
        company_patterns = [
            r'\b([A-Z][a-zA-Z\s&]+(?:Inc|LLC|Corp|Company|Ltd|Group|Solutions|Technologies|Systems)\.?)\b',
            r'\b([A-Z][a-zA-Z\s&]{2,20})\s+(?:-|–|,|\|)',  # Company names followed by separator
        ]
        
        for pattern in company_patterns:
            matches = re.findall(pattern, text)
            experience['companies'].extend(matches)
        
        # Calculate approximate years of experience
        years = re.findall(r'\b(19|20)\d{2}\b', text)
        if len(years) >= 2:
            years = [int(''.join(year)) for year in years]
            experience['total_years'] = max(years) - min(years)
        
        return experience
    
    def process_resume(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Process resume and extract all information"""
        # Extract raw text
        raw_text = self.extract_text(file_path, file_type)
        
        # Clean text
        processed_text = self.clean_text(raw_text)
        
        # Extract structured information
        contact_info = self.extract_contact_info(processed_text)
        education = self.extract_education(processed_text)
        experience = self.extract_experience(processed_text)
        
        return {
            'raw_text': raw_text,
            'processed_text': processed_text,
            'contact_info': contact_info,
            'education': education,
            'experience': experience
        }

# Create global instance
resume_processor = ResumeProcessor()